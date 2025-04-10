# word-limit.py: counts how many words a .docx or .pdf file has. 
# This can be used on a file for ACM CHI conference or ACM UIST, you can skip automatically over captions or include them, or even see the count with or without references. 
# How to use (in your command line): 
#   >python3 word-limit.py -f -v file.docx
#   ...this example counts words in a file, including captions of figures. Excludes references. 
#   >python3 word-limit.py --help
#   ... this will allow you to get started and see the options

# File type? This only supports .pdf or .docx 
# No flags are needed, the filetype is detected from extension

from docx import Document
import sys
import re
import argparse
import datetime
import subprocess

# figure is detected from the following template
# "Figure <number>. <text until end of line>"
#   So make sure to use "F" and not "f"
#   And use the period (.) as well after the <number>.
fig = re.compile(r"^Figure\s\d\.\s.*$")
abstract_seen = False

# limits (as of UIST 2025)
LIMIT_SHORT = 5000
LIMIT_AVERAGE = 8000
LIMIT_LONG = 12000

# words used to delimite sections
REFERENCES = "REFERENCES"
ABSTRACT = "ABSTRACT"
INTRODUCTION = "INTRODUCTION"

parser = argparse.ArgumentParser(description='Word-counter for papers in .docx format.')
parser.add_argument('-f', '--count_figure', action='store_true', help='Counts words inside of Figure captions')
parser.add_argument('-na', '--no_abstract', action='store_true', help='Does not count Authors, Fig 1 (if it comes before abstract), and abstract, i.e., starts at introduction')
parser.add_argument('-refs', '--count_references', action='store_true', help='Counts words in the reference section')
parser.add_argument('-v', '--verbose', action='store_true', help='verbose output indicates where words are used (e.g., Figures, Abstract, References, etc.')
parser.add_argument('-d', '--debug', action='store_true', help='Enables debug output that indicates when elements are skipped and not counted.')
parser.add_argument('-n', '--no_determination', action='store_true', help='Does not check the lenght of your paper against the limits set in the file (e.g., SHORT paper vs. LONG paper')
parser.add_argument('filename')

args = parser.parse_args()
paper = None

class Text: #honestly, kind of a ugly solution
    def __init__(self, pdf=None, doc=None):
        self.pdf = pdf
        self.doc = doc
    def paragraphs(self):
        if self.pdf is not None:
            return self.pdf
        else:
            return self.doc.paragraphs
    def text(self, par):
         print(par)
         if self.pdf is not None:
            return par
         else:
            return par.text

if args.filename[-4:].lower() == ".pdf":
    #print("PDF")
    now = datetime.datetime.now()
    formatted_date = now.strftime("%Y_%m_%d_%H_%M_%S")
    result = subprocess.run(['pdftotext', args.filename, 'out_' + formatted_date + '.txt'], capture_output=True, text=True)
    # note that pdftotext was run without layout flag
    with open('out_' + formatted_date + '.txt', 'r') as file:
        paper = Text(pdf=file.readlines())
elif args.filename[-5:].lower() == ".docx":
    #print("DOCX")
    paper = Text(doc=Document(args.filename))

# store counts
count_all = [] 
count_figures = []
count_references = []
REFERENCES_FOUND = False

for par in paper.paragraphs(): #no () 

    if fig.match(paper.text(par).strip()) is not None: # target text is a figure
        if args.count_figure: # and we have figure flag enabled
            count_all.append(paper.text(par)) # thus, counted figure
            count_figures.append(paper.text(par)) # counted figures for stats
        else: 
            count_figures.append(paper.text(par)) # counted figures for stats
            if args.debug:
                print("Debug\tFIGURE_SKIP\t" + paper.text(par))

    elif paper.text(par).strip() == REFERENCES: #passing over the references
        REFERENCES_FOUND = True
        continue

    elif args.count_references and REFERENCES_FOUND:
        count_all.append(par.text) # enabled so counted reference as main
        count_references.append(par.text) # count reference for stats
    elif not args.count_references and REFERENCES_FOUND:
        count_references.append(paper.text(par)) # count reference for stats
        if args.debug:
            print("Debug\tREFERENCE_SKIP\t" + paper.text(par))
    
    else: # catch all
        count_all.append(paper.text(par))
       
if args.debug: 
    print(len('\n'.join(count_all).split()))


if args.verbose:
    print("Breakdown")
    print("\ttotal of words in this run\t" + str(len('\n'.join(count_all).split())))
    print("\twords in figure caption text\t" + str(len('\n'.join(count_figures).split())))
    print("\twords in reference section\t" + str(len('\n'.join(count_references).split())))

if not args.no_determination:
    word_count = len('\n'.join(count_all).split())
    if word_count < LIMIT_SHORT:
       print("SHORT OK: " + str(word_count) + " is below " + str(LIMIT_SHORT)) 
    elif (word_count > LIMIT_SHORT and word_count < LIMIT_AVERAGE):
        print("NOT SHORT, AVERAGE OK: " + str(word_count) + " is below " + str(LIMIT_AVERAGE))
    elif (word_count > LIMIT_AVERAGE and word_count < LIMIT_LONG):
        print("NOT AVERAGE!, LONG OK: " + str(word_count) + " is below " + str(LIMIT_LONG))
    else: 
        print("VERY LONG!: " + str(word_count) + " is OVER " + str(LIMIT_LONG))
